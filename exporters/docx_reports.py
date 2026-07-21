"""Accessible Word renderers for strategy and approved content assets."""

from __future__ import annotations

import contextlib
import re
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from .brand import COPPER, INDIGO, INK, MUTED, PAPER, WHITE

# Word cannot embed fonts via python-docx, so the document must only name
# typefaces that ship with every Word install. Georgia/Calibri are the closest
# universally-available matches to the brand pair (Fraunces / Source Sans 3).
DOCX_DISPLAY_FONT = "Georgia"
DOCX_BODY_FONT = "Calibri"

_THEME_FONT_ATTRS = ("w:asciiTheme", "w:hAnsiTheme", "w:eastAsiaTheme", "w:cstheme")
_RFONT_SLOTS = ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs")


def _display(value: Any) -> str:
    """Render a measured value, or state plainly that it was not measured."""

    if value is None or value == "":
        return "Unavailable"
    return str(value)


def _rgb(value: str) -> RGBColor:
    value = value.lstrip("#")
    return RGBColor(int(value[0:2], 16), int(value[2:4], 16), int(value[4:6], 16))


def _pin_style_font(style: Any, font_name: str) -> None:
    """Pin every rFonts slot on a style and strip theme references.

    Word resolves ``asciiTheme``/``hAnsiTheme`` attributes ahead of explicit
    ``rFonts`` names, so both must be handled: set all four concrete slots and
    delete any theme attributes left over from the built-in style definitions.
    """

    style.font.name = font_name
    rfonts = style.element.get_or_add_rPr().get_or_add_rFonts()
    for slot in _RFONT_SLOTS:
        rfonts.set(qn(slot), font_name)
    for attr in _THEME_FONT_ATTRS:
        rfonts.attrib.pop(qn(attr), None)


def _pin_document_fonts(doc: Document) -> None:
    """Rewrite the theme part so theme-font fallbacks resolve to installed faces."""

    with contextlib.suppress(Exception):  # the theme part may be absent entirely
        theme = doc.part.part_related_by(RT.THEME)
        blob = theme.blob.decode("utf-8")
        blob = re.sub(
            r'(<a:majorFont>\s*<a:latin\b[^>]*?typeface=")[^"]*(")',
            lambda match: match.group(1) + DOCX_DISPLAY_FONT + match.group(2),
            blob,
        )
        blob = re.sub(
            r'(<a:minorFont>\s*<a:latin\b[^>]*?typeface=")[^"]*(")',
            lambda match: match.group(1) + DOCX_BODY_FONT + match.group(2),
            blob,
        )
        theme._blob = blob.encode("utf-8")


def _shade(cell: Any, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill.lstrip("#"))


def _cell_margins(cell: Any, top: int = 90, start: int = 105, bottom: int = 90, end: int = 105) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        tag = tc_mar.find(qn(f"w:{side}"))
        if tag is None:
            tag = OxmlElement(f"w:{side}")
            tc_mar.append(tag)
        tag.set(qn("w:w"), str(value))
        tag.set(qn("w:type"), "dxa")


def _repeat_header(row: Any) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    marker = OxmlElement("w:tblHeader")
    marker.set(qn("w:val"), "true")
    tr_pr.append(marker)


def _page_number(paragraph: Any) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for item in (begin, instruction, separate, text, end):
        run._r.append(item)


class DOCXReportBuilder:
    """Narrative-proposal preset with an editorial cover and explicit geometry."""

    def __init__(self, project_root: Path) -> None:
        self.project_root = project_root

    def _base_document(self, *, title: str, subject: str) -> Document:
        doc = Document()
        section = doc.sections[0]
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(0.85)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.9)
        section.right_margin = Inches(0.9)
        section.header_distance = Inches(0.3)
        section.footer_distance = Inches(0.35)

        core = doc.core_properties
        core.title = title
        core.subject = subject
        core.author = "Traffic Radius"
        core.keywords = "enterprise SEO, evidence, strategy, approval"
        core.comments = "Generated from canonical, evidence-linked records."

        styles = doc.styles
        normal = styles["Normal"]
        _pin_style_font(normal, DOCX_BODY_FONT)
        normal.font.size = Pt(10.5)
        normal.font.color.rgb = _rgb(INK)
        normal_rpr = normal._element.get_or_add_rPr()
        normal_lang = normal_rpr.find(qn("w:lang"))
        if normal_lang is None:
            normal_lang = OxmlElement("w:lang")
            normal_rpr.append(normal_lang)
        normal_lang.set(qn("w:val"), "en-AU")
        normal.paragraph_format.space_after = Pt(6)
        normal.paragraph_format.line_spacing = 1.22

        for name, size, color, before, after in (
            ("Title", 30, INK, 0, 12),
            ("Heading 1", 18, INK, 18, 8),
            ("Heading 2", 13, INDIGO, 14, 5),
            ("Heading 3", 11, COPPER, 10, 4),
        ):
            style = styles[name]
            _pin_style_font(
                style,
                DOCX_DISPLAY_FONT if name in {"Title", "Heading 1"} else DOCX_BODY_FONT,
            )
            style.font.size = Pt(size)
            style.font.color.rgb = _rgb(color)
            style.font.bold = name != "Title"
            style.paragraph_format.space_before = Pt(before)
            style.paragraph_format.space_after = Pt(after)
            style.paragraph_format.keep_with_next = True

        if "Evidence Callout" not in styles:
            callout = styles.add_style("Evidence Callout", WD_STYLE_TYPE.PARAGRAPH)
        else:
            callout = styles["Evidence Callout"]
        _pin_style_font(callout, DOCX_BODY_FONT)
        callout.font.size = Pt(11)
        callout.font.color.rgb = _rgb(INDIGO)
        callout.paragraph_format.left_indent = Inches(0.2)
        callout.paragraph_format.right_indent = Inches(0.2)
        callout.paragraph_format.space_before = Pt(8)
        callout.paragraph_format.space_after = Pt(10)

        header = section.header
        header.is_linked_to_previous = False
        header_table = header.add_table(rows=1, cols=2, width=Inches(6.7))
        header_table.columns[0].width = Inches(3.8)
        header_table.columns[1].width = Inches(2.9)
        left, right = header_table.rows[0].cells
        left.text = "TRAFFIC RADIUS"
        right.text = "ENTERPRISE SEO STUDIO"
        right.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        for cell in (left, right):
            _cell_margins(cell, top=0, bottom=40, start=0, end=0)
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(7.5)
                run.font.bold = True
                run.font.color.rgb = _rgb(MUTED)

        footer = section.footer
        footer.is_linked_to_previous = False
        footer_table = footer.add_table(rows=1, cols=2, width=Inches(6.7))
        footer_table.columns[0].width = Inches(5.6)
        footer_table.columns[1].width = Inches(1.1)
        footer_table.cell(0, 0).text = "Evidence before assertion · Private client document"
        page_paragraph = footer_table.cell(0, 1).paragraphs[0]
        page_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        _page_number(page_paragraph)
        for cell in footer_table.rows[0].cells:
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(7.5)
                run.font.color.rgb = _rgb(MUTED)

        _pin_document_fonts(doc)
        return doc

    def _cover(
        self,
        doc: Document,
        *,
        title: str,
        subtitle: str,
        client: str,
        as_of: str,
        run_id: str,
        status: str = "APPROVAL-READY",
    ) -> None:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(42)
        p.paragraph_format.space_after = Pt(18)
        run = p.add_run("TRAFFIC RADIUS · EVIDENCE-LED DELIVERY")
        run.font.size = Pt(9)
        run.font.bold = True
        run.font.color.rgb = _rgb(COPPER)

        doc.add_heading(title, 0)
        p = doc.add_paragraph(subtitle)
        p.paragraph_format.space_after = Pt(28)
        for run in p.runs:
            run.font.size = Pt(14)
            run.font.color.rgb = _rgb(MUTED)

        table = doc.add_table(rows=4, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        table.autofit = False
        labels = ("CLIENT", "EVIDENCE AS OF", "RUN", "STATUS")
        values = (client, as_of, run_id, status)
        for index, (label, value) in enumerate(zip(labels, values, strict=True)):
            table.cell(index, 0).width = Inches(1.35)
            table.cell(index, 1).width = Inches(5.15)
            table.cell(index, 0).text = label
            table.cell(index, 1).text = value
            for cell in table.rows[index].cells:
                _cell_margins(cell)
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for run in table.cell(index, 0).paragraphs[0].runs:
                run.font.size = Pt(8)
                run.font.bold = True
                run.font.color.rgb = _rgb(MUTED)
            for run in table.cell(index, 1).paragraphs[0].runs:
                run.font.size = Pt(9)
                run.font.color.rgb = _rgb(INK)

        p = doc.add_paragraph(style="Evidence Callout")
        p.paragraph_format.space_before = Pt(34)
        p.add_run(
            "Measured and derived claims in this document resolve to the evidence register. Unavailable private data is named explicitly; no substitute metrics were invented."
        )
        doc.add_page_break()

    def _table(self, doc: Document, headers: list[str], rows: list[list[Any]], widths: list[float]) -> Any:
        table = doc.add_table(rows=1, cols=len(headers))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        for grid_column, width in zip(table._tbl.tblGrid.gridCol_lst, widths, strict=True):
            grid_column.w = Inches(width)
        table.style = "Table Grid"
        for index, (header, width) in enumerate(zip(headers, widths, strict=True)):
            cell = table.rows[0].cells[index]
            cell.width = Inches(width)
            cell.text = header
            _shade(cell, INDIGO)
            _cell_margins(cell)
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(8)
                run.font.bold = True
                run.font.color.rgb = _rgb(WHITE)
        _repeat_header(table.rows[0])
        for row_index, values in enumerate(rows):
            cells = table.add_row().cells
            for index, (value, width) in enumerate(zip(values, widths, strict=True)):
                cells[index].width = Inches(width)
                cells[index].text = str(value) if value not in (None, "") else "Unavailable"
                cells[index].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
                _cell_margins(cells[index])
                if row_index % 2:
                    _shade(cells[index], PAPER)
                for run in cells[index].paragraphs[0].runs:
                    run.font.size = Pt(8)
                    run.font.color.rgb = _rgb(INK)
        doc.add_paragraph().paragraph_format.space_after = Pt(0)
        return table

    def strategy_report(self, data: dict[str, Any], output: Path) -> Path:
        output.parent.mkdir(parents=True, exist_ok=True)
        run = data["run"]
        client_name = data["client"]["name"]
        doc = self._base_document(
            title=f"{client_name} Enterprise SEO Strategy",
            subject="Evidence-led 16-week enterprise SEO strategy",
        )
        self._cover(
            doc,
            title="Enterprise SEO Strategy",
            subtitle=f"A disciplined, evidence-linked operating strategy for {client_name}.",
            client=data["client"]["name"],
            as_of=run["evidence_as_of"],
            run_id=run["id"],
        )
        doc.add_heading("Executive direction", level=1)
        doc.add_paragraph(data["executive_summary"])
        callout = doc.add_paragraph(style="Evidence Callout")
        callout.add_run(run["overall_score_reason"])

        for section in data.get("strategy_sections", []):
            level = min(max(int(section.get("level", 1)), 1), 3)
            doc.add_heading(section["title"], level=level)
            for paragraph in section.get("paragraphs", []):
                doc.add_paragraph(paragraph)
            if section.get("decision"):
                callout = doc.add_paragraph(style="Evidence Callout")
                marker = callout.add_run("Strategic decision · ")
                marker.bold = True
                callout.add_run(section["decision"])

        doc.add_heading("Canonical opportunity map", level=1)
        self._table(
            doc,
            ["Cluster", "Intent", "Target URL", "Decision", "Evidence"],
            [
                [
                    item["cluster"],
                    item["intent"],
                    item["target_url"],
                    item["decision"],
                    ", ".join(item["evidence_ids"]),
                ]
                for item in data.get("opportunities", [])
            ],
            [1.25, 0.8, 1.8, 1.35, 0.75],
        )

        doc.add_heading("Measurement contract", level=1)
        doc.add_paragraph(
            "Baselines remain unavailable until the named private source is connected. The team must capture and approve the baseline before setting a scenario band or target."
        )
        self._table(
            doc,
            ["KPI", "Baseline", "Cadence", "Source", "Decision use"],
            [
                [item["kpi"], item["baseline"], item["cadence"], item["source"], item["decision_use"]]
                for item in data.get("measurement_plan", [])
            ],
            [1.2, 1.0, 0.8, 1.2, 2.0],
        )

        doc.add_heading("Evidence and limitations", level=1)
        self._table(
            doc,
            ["ID", "Source", "Captured", "Scope", "Status"],
            [
                [source["id"], source["label"], source["captured_at"], source["scope"], source["status"]]
                for source in data.get("sources", [])
            ],
            [0.65, 2.1, 1.05, 1.55, 0.65],
        )
        for limitation in data.get("limitations", []):
            doc.add_paragraph(limitation, style="List Bullet")
        doc.save(output)
        return output

    def content_strategy(self, data: dict[str, Any], output: Path) -> Path:
        """Render the content-and-keyword strategy as its own deliverable.

        Deliberately distinct from ``strategy_report``: this document answers
        "what should we publish and for which demand", using the keyword,
        cluster and content-asset evidence rather than the operating plan.
        """

        output.parent.mkdir(parents=True, exist_ok=True)
        run = data["run"]
        client_name = data["client"]["name"]
        doc = self._base_document(
            title=f"{client_name} Content and Keyword Strategy",
            subject="Evidence-led content strategy and demand mapping",
        )
        self._cover(
            doc,
            title="Content and Keyword Strategy",
            subtitle=f"Demand, coverage and publishing priorities for {client_name}.",
            client=client_name,
            as_of=run["evidence_as_of"],
            run_id=run["id"],
        )

        market = data.get("market") or {}
        keywords = list(data.get("keywords") or [])
        clusters = list(data.get("keyword_clusters") or [])
        assets = list(data.get("content_assets") or [])

        doc.add_heading("Demand evidence", level=1)
        if market.get("status") == "available":
            domain = market.get("domain") or {}
            doc.add_paragraph(
                f"{len(keywords)} ranking keywords were retrieved from "
                f"{market.get('provider', 'the connected provider')} for the "
                f"{market.get('database', 'configured')} database. The domain holds "
                f"{_display(domain.get('organic_keywords'))} organic keywords and "
                f"{_display(domain.get('organic_traffic'))} estimated monthly organic sessions."
            )
        else:
            doc.add_paragraph(
                "Search-demand metrics are unavailable for this run: "
                f"{market.get('unavailable_reason') or 'no keyword provider is connected'}. "
                "Every priority below is therefore derived from crawl evidence only, and no "
                "volume, difficulty or traffic figure is asserted."
            )

        doc.add_heading("Topic clusters and coverage", level=1)
        self._table(
            doc,
            ["Cluster", "Keywords", "Volume", "Coverage", "Primary URL"],
            [
                [
                    cluster.get("name", "Unavailable"),
                    cluster.get("keyword_count", 0),
                    _display(cluster.get("total_volume")),
                    cluster.get("coverage", "unknown"),
                    cluster.get("primary_url") or "No mapped page",
                ]
                for cluster in clusters[:40]
            ]
            or [["No clusters were derived from the available evidence.", "", "", "", ""]],
            [1.6, 0.75, 0.8, 1.0, 2.0],
        )

        doc.add_heading("Priority keywords", level=1)
        self._table(
            doc,
            ["Keyword", "Position", "Volume", "CPC", "Landing URL"],
            [
                [
                    item.get("phrase", ""),
                    _display(item.get("position")),
                    _display(item.get("search_volume")),
                    _display(item.get("cpc")),
                    item.get("landing_url") or "Unmapped",
                ]
                for item in keywords[:40]
            ]
            or [["No provider keyword data was available for this run.", "", "", "", ""]],
            [2.0, 0.75, 0.8, 0.7, 1.9],
        )

        doc.add_heading("Publishing priorities", level=1)
        doc.add_paragraph(
            "Each asset below exists because a crawled page or an evidenced demand cluster "
            "justified it. Drafts stay withheld until a human approves them."
        )
        self._table(
            doc,
            ["Asset", "Intent", "Target URL", "Approval"],
            [
                [
                    asset.get("title", ""),
                    asset.get("intent", ""),
                    asset.get("target_url", ""),
                    asset.get("approval_state", ""),
                ]
                for asset in assets
            ]
            or [["No content assets passed the evidence gates for this run.", "", "", ""]],
            [2.1, 1.1, 1.85, 1.1],
        )

        for limitation in data.get("limitations", []):
            doc.add_paragraph(limitation, style="List Bullet")
        doc.save(output)
        return output

    def content_asset(self, data: dict[str, Any], asset: dict[str, Any], output: Path) -> Path:
        output.parent.mkdir(parents=True, exist_ok=True)
        run = data["run"]
        doc = self._base_document(title=asset["title"], subject="Evidence-supported SEO content asset")
        self._cover(
            doc,
            title=asset["title"],
            subtitle=f"{asset['asset_type']} · Target: {asset['target_url']}",
            client=data["client"]["name"],
            as_of=run["evidence_as_of"],
            run_id=run["id"],
            status=asset["approval_state"].upper(),
        )
        doc.add_heading("Editorial brief", level=1)
        self._table(
            doc,
            ["Audience", "Intent", "Primary topic", "Target URL"],
            [[asset["audience"], asset["intent"], asset["primary_topic"], asset["target_url"]]],
            [1.5, 1.15, 1.7, 1.95],
        )
        doc.add_heading(asset["headline"], level=1)
        for block in asset.get("body", []):
            if block["type"] == "heading":
                doc.add_heading(block["text"], level=min(int(block.get("level", 2)), 3))
            elif block["type"] == "list":
                for item in block["items"]:
                    doc.add_paragraph(item, style="List Bullet")
            else:
                doc.add_paragraph(block["text"])

        doc.add_page_break()
        doc.add_heading("Claim ledger", level=1)
        self._table(
            doc,
            ["Claim", "Evidence", "Confidence", "Validation"],
            [
                [
                    claim["claim"],
                    ", ".join(claim["evidence_ids"]),
                    f"{claim['confidence']:.0%}",
                    claim["validation"],
                ]
                for claim in asset.get("claims", [])
            ],
            [2.7, 1.0, 0.8, 1.8],
        )
        doc.add_heading("Source ledger", level=1)
        source_ids = {item for claim in asset.get("claims", []) for item in claim["evidence_ids"]}
        self._table(
            doc,
            ["ID", "Source", "Captured", "Scope"],
            [
                [source["id"], source["label"], source["captured_at"], source["scope"]]
                for source in data.get("sources", [])
                if source["id"] in source_ids
            ],
            [0.7, 2.5, 1.15, 2.0],
        )
        doc.save(output)
        return output

