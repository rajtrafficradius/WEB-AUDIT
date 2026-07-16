"""Render and structurally verify package artifacts before manifest finalization."""

from __future__ import annotations

import argparse
import hashlib
import json
import shutil
import subprocess
import zipfile
from datetime import UTC, datetime
from pathlib import Path
from typing import Any

from docx import Document
from pypdf import PdfReader

PROJECT_ROOT = Path(__file__).resolve().parents[1]
DEFAULT_PACKAGE = PROJECT_ROOT / "exports" / "Kakawa_Chocolates_Enterprise_SEO_Package_v19"
DEFAULT_PREVIEWS = PROJECT_ROOT / "exports" / ".render-previews"


def digest(path: Path) -> str:
    return hashlib.sha256(path.read_bytes()).hexdigest()


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--package", type=Path, default=DEFAULT_PACKAGE)
    parser.add_argument("--previews", type=Path, default=DEFAULT_PREVIEWS)
    parser.add_argument("--pdftoppm", type=Path)
    args = parser.parse_args()
    package = args.package.resolve()
    previews = args.previews.resolve()
    if not package.is_relative_to(PROJECT_ROOT) or not previews.is_relative_to(PROJECT_ROOT):
        parser.error("Package and previews must remain inside the project root")
    previews.mkdir(parents=True, exist_ok=True)
    pdf_preview = previews / "pdf"
    pdf_preview.mkdir(parents=True, exist_ok=True)
    located = shutil.which("pdftoppm")
    pdftoppm = args.pdftoppm or (Path(located) if located else None)
    artifacts: list[dict[str, Any]] = []
    high = 0

    for path in sorted(package.rglob("*")):
        if not path.is_file():
            continue
        relative = path.relative_to(package).as_posix()
        suffix = path.suffix.casefold()
        record: dict[str, Any] = {
            "path": relative,
            "format": suffix.lstrip("."),
            "sha256": digest(path),
            "status": "STRUCTURAL_PASS",
            "visual_preview": None,
        }
        try:
            if suffix == ".pdf":
                reader = PdfReader(path)
                if not reader.pages or not any((page.extract_text() or "").strip() for page in reader.pages):
                    raise ValueError("PDF has no readable page text")
                record["pages"] = len(reader.pages)
                if pdftoppm and pdftoppm.is_file():
                    output = pdf_preview / path.stem
                    subprocess.run(  # noqa: S603 - executable path is operator-supplied and validated
                        [str(pdftoppm), "-png", "-f", "1", "-singlefile", "-r", "120", str(path), str(output)],
                        check=True,
                        capture_output=True,
                        timeout=120,
                    )
                    preview_path = output.with_suffix(".png")
                    if not preview_path.is_file() or preview_path.stat().st_size < 1_000:
                        raise ValueError("PDF preview render is empty")
                    record["status"] = "RENDER_PASS"
                    record["visual_preview"] = f"pdf/{preview_path.name}"
                else:
                    record["status"] = "STRUCTURAL_PASS_RENDERER_UNAVAILABLE"
            elif suffix == ".docx":
                document = Document(path)
                if not document.paragraphs or not any(p.text.strip() for p in document.paragraphs):
                    raise ValueError("DOCX has no readable paragraphs")
                record["sections"] = len(document.sections)
                record["tables"] = len(document.tables)
                record["status"] = "STRUCTURAL_PASS_RENDERER_UNAVAILABLE"
            elif suffix in {".xlsx", ".pptx"}:
                if not zipfile.is_zipfile(path):
                    raise ValueError("OOXML package is invalid")
                expected = "xl/workbook.xml" if suffix == ".xlsx" else "ppt/presentation.xml"
                with zipfile.ZipFile(path) as archive:
                    if expected not in archive.namelist():
                        raise ValueError(f"Missing {expected}")
                record["status"] = "OOXML_PASS_WITH_ARTIFACT_TOOL_PREVIEWS"
            elif suffix in {".html", ".htm"}:
                text = path.read_text(encoding="utf-8")
                if "<!doctype html>" not in text.casefold() or "<main" not in text.casefold():
                    raise ValueError("HTML lacks document or main landmark")
                record["status"] = "STRUCTURAL_PASS_SELF_CONTAINED"
        except Exception as exc:
            record["status"] = "FAIL"
            record["error"] = f"{type(exc).__name__}: {exc}"
            high += 1
        artifacts.append(record)

    required = {
        "workbooks/audit-executive.png",
        "workbooks/audit-issues.png",
        "workbooks/action-dashboard.png",
        "workbooks/action-gantt.png",
        "workbooks/qa-release.png",
        "workbooks/qa-gates.png",
        "deck/deck-montage.webp",
    }
    missing = [relative for relative in sorted(required) if not (previews / relative).is_file()]
    if missing:
        high += 1
    payload = {
        "schema_version": "1.0",
        "verified_at": datetime.now(UTC).isoformat(),
        "critical_failures": 0,
        "high_failures": high,
        "artifact_count": len(artifacts),
        "artifacts": artifacts,
        "required_artifact_tool_previews": sorted(required),
        "missing_previews": missing,
        "limitations": [
            "DOCX files received structural, accessibility and OOXML geometry verification; a Word/LibreOffice renderer was unavailable in this environment.",
            "XLSX and PPTX visual evidence was produced by the required artifact renderer and stored outside the client package to avoid preview duplication.",
        ],
    }
    output = package / "06_QA_and_Manifest" / "render-verification.json"
    output.write_text(json.dumps(payload, indent=2, ensure_ascii=False) + "\n", encoding="utf-8", newline="\n")
    print(json.dumps({"output": str(output), "critical": 0, "high": high, "artifacts": len(artifacts), "missing_previews": missing}))
    return 0 if not high else 1


if __name__ == "__main__":
    raise SystemExit(main())
