from __future__ import annotations

import re
from pathlib import Path
from typing import Any

import pytest
from docx import Document
from docx.oxml.ns import qn
from pypdf import PdfReader

from exporters.docx_reports import DOCXReportBuilder
from exporters.html_outputs import build_content_html, build_html_deck
from exporters.pdf_reports import PDFReportBuilder

MACHINE_PATH = re.compile(r"(?:file:/+|(?<![A-Za-z0-9])[A-Za-z]:[\\/])", re.IGNORECASE)
MOJIBAKE_MARKERS = ("\ufffd", "\u00c2\u00b7", "\u00e2\u20ac", "\u00c3")


def _sample_data() -> dict[str, Any]:
    return {
        "client": {
            "name": "Kakawa Chocolates",
            "domain": "kakawachocolates.com.au",
        },
        "run": {
            "id": "RUN-KAKAWA-V19-TEST",
            "evidence_as_of": "2026-07-15",
            "evidence_coverage": 0.82,
            "coverage_interpretation": "Enough weighted evidence to publish category scores.",
            "overall_score": 76,
            "overall_score_reason": "Published because weighted evidence coverage exceeds 70%.",
        },
        "executive_summary": (
            "Observed crawl evidence supports a controlled technical clean-up before expansion."
        ),
        "categories": [
            {"category": "Technical", "score": 72},
            {"category": "Content", "score": 81},
        ],
        "findings": [
            {
                "priority": "P1",
                "title": "Consolidate a duplicate collection route",
                "impact": "Reduces competing indexable routes for the same intent.",
                "confidence": 0.94,
                "evidence_ids": ["EV-001"],
            }
        ],
        "actions": [
            {
                "id": "ACT-001",
                "week": 1,
                "week_end": 2,
                "priority": "P1",
                "action": "Validate and approve the canonical route.",
                "owner": "SEO lead",
                "approval_class": "admin_required",
                "kpi": "One approved canonical decision",
                "dependencies": [],
                "effort": "M",
            }
        ],
        "sources": [
            {
                "id": "SRC-001",
                "label": "Approved-domain crawl",
                "captured_at": "2026-07-15T05:30:00Z",
                "scope": "Public HTML pages on kakawachocolates.com.au",
                "status": "available",
            }
        ],
        "limitations": [
            "Private analytics remain unavailable until an authorised connection is approved."
        ],
        "strategy_sections": [
            {
                "title": "Technical integrity first",
                "level": 1,
                "paragraphs": [
                    "Resolve indexation ambiguity before creating additional landing pages."
                ],
                "decision": "Require an administrator to approve every canonical change.",
            },
            {
                "title": "Evidence-led expansion",
                "level": 2,
                "paragraphs": [
                    "Create content only when a distinct opportunity survives support "
                    "and overlap checks."
                ],
            },
        ],
        "opportunities": [
            {
                "cluster": "Chocolate gifts",
                "intent": "Commercial",
                "target_url": "https://kakawachocolates.com.au/collections/gifts",
                "decision": "Improve existing page",
                "evidence_ids": ["EV-001"],
            }
        ],
        "measurement_plan": [
            {
                "kpi": "Valid canonical coverage",
                "baseline": "Capture after approval",
                "cadence": "Weekly",
                "source": "Approved-domain crawl",
                "decision_use": "Confirm consolidation is safe.",
            }
        ],
        "qa": {
            "release_status": "PASS",
            "release_statement": "No unresolved Critical or High release failures.",
            "gates": [
                {
                    "name": "Domain safety",
                    "status": "PASS",
                    "critical_failures": 0,
                    "high_failures": 0,
                    "evidence": "Approved-domain scan",
                }
            ],
            "reconciliation": [
                {
                    "measure": "Findings",
                    "canonical": 1,
                    "package": 1,
                    "result": "PASS",
                }
            ],
        },
        "comparison": [],
        "deck": [
            {
                "eyebrow": "Gate 1 decision",
                "title": "Protect evidence before expansion",
                "body": "Approve the technical direction before generating deployment assets.",
                "points": [
                    {"label": "Coverage", "text": "82% weighted evidence coverage"},
                    {"label": "External changes", "text": "None"},
                ],
            },
            {
                "eyebrow": "First move",
                "title": "Resolve the canonical route",
                "body": "One approval-controlled decision removes a material ambiguity.",
                "points": [{"label": "Evidence", "text": "EV-001"}],
            },
        ],
    }


def _sample_asset() -> dict[str, Any]:
    return {
        "id": "CNT-001",
        "title": "Chocolate gift guide",
        "asset_type": "Editorial guide",
        "headline": "A considered guide to chocolate gifts",
        "summary": "Choose a format that matches the occasion and recipient.",
        "target_url": "https://kakawachocolates.com.au/collections/gifts",
        "audience": "Australian gift buyers",
        "intent": "Commercial investigation",
        "primary_topic": "Chocolate gifts",
        "approval_state": "approved",
        "body": [
            {"type": "heading", "level": 2, "text": "Start with the occasion"},
            {"type": "paragraph", "text": "Delivery timing and format shape the choice."},
            {"type": "list", "items": ["Confirm timing", "Review storage guidance"]},
        ],
        "claims": [
            {
                "claim": "The gifts collection was observed on the approved domain.",
                "evidence_ids": ["SRC-001"],
                "confidence": 1.0,
                "validation": "domain_and_link_check_passed",
            }
        ],
    }


def _assert_clean_text(text: str) -> None:
    assert not MACHINE_PATH.search(text), "A machine-specific path leaked into a client artifact"
    for marker in MOJIBAKE_MARKERS:
        assert marker not in text


def _table_grid_widths(table: Any) -> list[int]:
    grid_columns = table._tbl.xpath("./w:tblGrid/w:gridCol")
    return [int(column.get(qn("w:w"))) for column in grid_columns]


def _cell_widths(row: Any) -> list[int]:
    widths: list[int] = []
    for cell in row.cells:
        width = cell._tc.find(qn("w:tcPr")).find(qn("w:tcW"))
        assert width is not None
        assert width.get(qn("w:type")) == "dxa"
        widths.append(int(width.get(qn("w:w"))))
    return widths


def test_html_deck_is_self_contained_escaped_and_path_independent(tmp_path: Path) -> None:
    data = _sample_data()
    data["deck"][0].update(
        {
            "eyebrow": '<img src=x onerror="alert(1)">',
            "title": "Evidence <script>alert(1)</script>",
            "body": 'A literal C-less marker: <b data-x="1">review</b>.',
            "points": [{"label": "<unsafe>", "text": '"quoted" & supported'}],
        }
    )

    output = build_html_deck(data, tmp_path / "Executive_Deck.html")
    rendered = output.read_text(encoding="utf-8")

    assert "&lt;script&gt;alert(1)&lt;/script&gt;" in rendered
    assert "&lt;img src=x onerror=&quot;alert(1)&quot;&gt;" in rendered
    assert "&quot;quoted&quot; &amp; supported" in rendered
    assert "<script" not in rendered.casefold()
    assert "<img" not in rendered.casefold()
    assert "<link" not in rendered.casefold()
    assert "url(http" not in rendered.casefold()
    assert "https://" not in rendered.casefold()
    assert 'lang="en-AU"' in rendered
    assert 'href="#slide-1"' in rendered
    _assert_clean_text(rendered)


def test_content_html_escapes_every_canonical_record_and_has_accessible_tables(
    tmp_path: Path,
) -> None:
    data = _sample_data()
    asset = _sample_asset()
    asset["headline"] = "Gifts <script>not executable</script>"
    asset["target_url"] = 'https://kakawachocolates.com.au/?q="gift"&x=<unsafe>'
    asset["body"][1]["text"] = "Use <em>literal source text</em> & verify it."
    asset["claims"][0]["claim"] = "Observed <b>source text</b> & nothing more."
    data["sources"][0]["label"] = "Crawler <primary> & replay"

    output = build_content_html(data, asset, tmp_path / "content.html")
    rendered = output.read_text(encoding="utf-8")

    assert "Gifts &lt;script&gt;not executable&lt;/script&gt;" in rendered
    assert "Use &lt;em&gt;literal source text&lt;/em&gt; &amp; verify it." in rendered
    assert "Observed &lt;b&gt;source text&lt;/b&gt; &amp; nothing more." in rendered
    assert "Crawler &lt;primary&gt; &amp; replay" in rendered
    assert "<script" not in rendered.casefold()
    assert rendered.count('<th scope="col">') == 8
    assert "<thead>" in rendered and "<tbody>" in rendered
    assert 'lang="en-AU"' in rendered
    _assert_clean_text(rendered)


@pytest.mark.render
def test_docx_strategy_has_semantic_structure_accessibility_and_exact_table_geometry(
    tmp_path: Path,
) -> None:
    output = DOCXReportBuilder(tmp_path).strategy_report(
        _sample_data(), tmp_path / "Kakawa_Enterprise_SEO_Strategy.docx"
    )
    document = Document(output)

    assert document.core_properties.title == "Kakawa Chocolates Enterprise SEO Strategy"
    assert document.core_properties.author == "Traffic Radius"
    assert document.core_properties.subject == "Evidence-led 16-week enterprise SEO strategy"

    normal_rpr = document.styles["Normal"]._element.get_or_add_rPr()
    language = normal_rpr.find(qn("w:lang"))
    assert language is not None, "The document default language must be explicit"
    assert language.get(qn("w:val")) == "en-AU"

    heading_levels = [
        int(paragraph.style.name.rsplit(" ", 1)[1])
        for paragraph in document.paragraphs
        if paragraph.style and paragraph.style.name.startswith("Heading ")
    ]
    assert heading_levels
    highest_seen = 0
    for level in heading_levels:
        assert level <= highest_seen + 1, (
            f"Heading hierarchy skipped from {highest_seen} to {level}"
        )
        highest_seen = max(highest_seen, level)

    bullet_text = {
        paragraph.text
        for paragraph in document.paragraphs
        if paragraph.style and paragraph.style.name == "List Bullet"
    }
    assert set(_sample_data()["limitations"]).issubset(bullet_text)

    assert len(document.tables) >= 4
    for table in document.tables[1:]:  # Cover metadata is a definition-list table.
        assert table.rows[0]._tr.xpath("./w:trPr/w:tblHeader"), (
            "Every data table needs a repeating semantic header row"
        )
        grid_widths = _table_grid_widths(table)
        assert len(grid_widths) == len(table.columns)
        for row in table.rows:
            assert _cell_widths(row) == grid_widths, (
                "tblGrid and every tcW must agree for deterministic Word geometry"
            )
            for cell in row.cells:
                margins = cell._tc.xpath("./w:tcPr/w:tcMar")
                assert margins, "Table cells require explicit breathing-room margins"

    footer_xml = "".join(footer._element.xml for footer in (s.footer for s in document.sections))
    assert " PAGE " in footer_xml
    assert 'w:fldCharType="begin"' in footer_xml
    assert 'w:fldCharType="end"' in footer_xml

    visible_text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    for table in document.tables:
        visible_text += "\n" + "\n".join(cell.text for row in table.rows for cell in row.cells)
    _assert_clean_text(visible_text)


@pytest.mark.render
def test_pdf_report_has_metadata_page_contract_and_literal_safe_text(tmp_path: Path) -> None:
    data = _sample_data()
    data["executive_summary"] = (
        "Treat <b>this evidence marker</b> as literal source text, not ReportLab markup."
    )
    output = PDFReportBuilder(tmp_path).executive_report(
        data, tmp_path / "Kakawa_Executive_Report.pdf"
    )
    reader = PdfReader(output)

    assert len(reader.pages) >= 3
    assert reader.metadata.title == "Executive SEO Review"
    assert reader.metadata.author == "Traffic Radius"
    extracted_pages = [(page.extract_text() or "") for page in reader.pages]
    combined = "\n".join(extracted_pages)
    first_page = " ".join(extracted_pages[0].split())

    assert "Enterprise SEO Evidence & Direction" in first_page
    assert "Kakawa Chocolates" in first_page
    assert "2026-07-15" in first_page
    assert "RUN-KAKAWA-V19-TEST" in first_page
    assert "<b>this evidence marker</b>" in combined
    assert "Priority findings" in combined
    assert "Source register" in combined
    assert all("TRAFFIC RADIUS" in page for page in extracted_pages)
    _assert_clean_text(combined)
